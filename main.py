from textwrap import dedent

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.table import _Cell
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.oxml.ns import nsdecls, qn


CHECK_BOX = u" □ "
# font functions
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_margins(cell: _Cell, **kwargs):
    """
    cell:  actual cell instance you want to modify

    usage:

        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)

def get_merge_cells(table,row,start,end):
    start_cell = table.cell(row, start)
    end_cell = table.cell(row, end)
    new_cell = start_cell.merge(end_cell)
    return new_cell

def set_table_alignment(table):
    for row in table.rows:
        row.height = Inches(0.25)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


# text functions
def get_header(header_text='RENTAL APPLICATION'):
    header = document.add_heading(header_text, 1)
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_font = header.style.font
    header_font.color.rgb = RGBColor(0, 0, 0)
    header_font.name = 'Calibri'
    header_font.size = Pt(20)


def get_header_paragraph(text):
    header_paragraph = document.add_paragraph()
    header_paragraph_text = f'''
    Property Address: {text}
    Unit #: {text}
    City, State, ZIP: {text}
    Date of Application: {text}
    '''
    header_paragraph_font = header_paragraph.style.font
    header_paragraph_font.name = 'Calibri'
    header_paragraph_font.size = Pt(12)
    header_paragraph.add_run(dedent(header_paragraph_text)).bold = True


def get_table_applicant_information():

    initials = ['First Name','','Middle Name','','Last Name','']
    contact_informations = ['Email','','Phone #1','','Phone #2','']
    document_informations = ['Date of Birth','_ _/ _ _/ _ _ _ _','Social Security #','','Driver’s License #','']

    table_applicant_information = document.add_table(rows=4,cols=6)

    table_applicant_information.autofit = False

    header_cell = get_merge_cells(table_applicant_information,0,0,5)
    header_cell.text = 'APPLICANT INFORMATION'
    table_header_color = parse_xml(r'<w:shd {} w:fill="#1f3864"/>'.format(nsdecls('w')))
    header_cell._tc.get_or_add_tcPr().append(table_header_color)

    for row in range(4):
        for cell in range(6):
            if cell % 2 != 0:
                set_cell_border(table_applicant_information.cell(row, cell), bottom={"sz": 6, "color": "#000000", "val": "single"})

    row_second_cells = table_applicant_information.rows[1].cells
    for cell, initial in enumerate(initials):
        row_second_cells[cell].text = initial

    row_third_cells = table_applicant_information.rows[2].cells
    for cell, contact_information in enumerate(contact_informations):
        row_third_cells[cell].text = contact_information

    row_fourth_cells = table_applicant_information.rows[3].cells

    for cell, document_information in enumerate(document_informations):
        row_fourth_cells[cell].text = document_information

    header_paragraph = document.add_paragraph()
    header_paragraph.aligmnet = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_paragraph_font = header_paragraph.style.font
    header_paragraph_font.name = 'Calibri'
    header_paragraph_font.size = Pt(12)


def get_table_additional_occupant(rows=4):
    table_additional_occupant = document.add_table(rows=rows, cols=7)

    table_additional_occupant.autofit = False

    paragraph = document.add_paragraph()
    paragraph.aligmnet = WD_TAB_ALIGNMENT
    paragraph_font = paragraph.style.font
    paragraph_font.name = 'Calibri'
    paragraph_font.size = Pt(10)


    header_cell = get_merge_cells(table_additional_occupant,0,0,6)
    header_cell.text = 'ADDITIONAL OCCUPANT(S)'
    table_header_color = parse_xml(r'<w:shd {} w:fill="#1f3864"/>'.format(nsdecls('w')))
    header_cell._tc.get_or_add_tcPr().append(table_header_color)


    for row in range(1, rows):
        column = table_additional_occupant.rows[row].cells
        column[0].text = 'Name'
        column[0].width = Inches(0.5)
        column_1_2 = get_merge_cells(table_additional_occupant, row, 1, 2)
        set_cell_border(column_1_2, bottom={"sz": 6, "color": "#000000", "val": "single"})

        column[3].text = 'Relationship'
        column[3].width = Inches(0.9)
        set_cell_border(column[4], bottom={"sz": 6, "color": "#000000", "val": "single"})

        column[5].text = 'Age'
        column[5].width = Inches(0.5)
        set_cell_border(column[6], bottom={"sz": 6, "color": "#000000", "val": "single"})


def get_table_residence_history():
    table_residence_history = document.add_table(rows=5, cols=14)

    for index, column in enumerate(table_residence_history.columns):
        for cell in column.cells:
            if index in [0, 1]:
                cell.width = Inches(0.75)
            elif index in [6, 7, 10, 11]:
                cell.width = Inches(0.6)
            else:
                cell.width = Inches(0.475)

    i = 0
    get_merge_cells(table_residence_history, i, 0, len(table_residence_history.columns) - 1)

    current_row = table_residence_history.rows[i]

    table_header_color = parse_xml(r'<w:shd {} w:fill="#1f3864"/>'.format(nsdecls('w')))
    current_row.cells[0]._tc.get_or_add_tcPr().append(table_header_color)

    current_row.cells[0].text = 'residence history'.upper()

    i = 1
    get_merge_cells(table_residence_history, i, 0, 1)
    get_merge_cells(table_residence_history, i, 2, 5)
    get_merge_cells(table_residence_history, i, 6, 7)
    get_merge_cells(table_residence_history, i, 8, 9)

    current_row = table_residence_history.rows[i]

    current_row.cells[0].text = 'Current Address'
    current_row.cells[6].text = 'Unit #'
    current_row.cells[10].text = 'Rent'
    current_row.cells[11].text = u"  □  "
    current_row.cells[12].text = 'Own'
    current_row.cells[13].text = u"  □  "

    set_cell_border(current_row.cells[2], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[8], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[11], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[13], bottom={"sz": 6, "color": "#000000", "val": "single"})

    i = 2
    get_merge_cells(table_residence_history, i, 6, 9)
    get_merge_cells(table_residence_history, i, 10, 13)

    current_row = table_residence_history.rows[i]

    current_row.cells[0].text = 'City'
    current_row.cells[2].text = 'State'
    current_row.cells[4].text = 'ZIP'
    current_row.cells[6].text = 'Monthly Payment or Rent $'

    set_cell_border(current_row.cells[1], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[3], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[5], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[10], bottom={"sz": 6, "color": "#000000", "val": "single"})

    i = 3
    get_merge_cells(table_residence_history, i, 0, 1)
    get_merge_cells(table_residence_history, i, 2, 5)
    get_merge_cells(table_residence_history, i, 6, 7)
    get_merge_cells(table_residence_history, i, 8, 9)
    get_merge_cells(table_residence_history, i, 10, 11)
    get_merge_cells(table_residence_history, i, 12, 13)

    current_row = table_residence_history.rows[i]

    current_row.cells[0].text = 'Dates of Residence'
    current_row.cells[2].text = f'__/__/____ to __/__/____'
    current_row.cells[6].text = 'Present Landlord'
    current_row.cells[10].text = 'Landlord Phone#'

    set_cell_border(current_row.cells[2], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[8], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[12], bottom={"sz": 6, "color": "#000000", "val": "single"})

    i = 4
    get_merge_cells(table_residence_history, i, 0, 1)
    get_merge_cells(table_residence_history, i, 2, len(table_residence_history.columns) - 1)

    current_row = table_residence_history.rows[i]

    current_row.cells[0].text = 'Reason for moving out'

    set_cell_border(current_row.cells[2], bottom={"sz": 6, "color": "#000000", "val": "single"})


def get_table_residence_history_no_header():
    table_residence_history = document.add_table(rows=4, cols=14)

    for index, column in enumerate(table_residence_history.columns):
        for cell in column.cells:
            if index in [0, 1]:
                cell.width = Inches(0.75)
            elif index in [6, 7, 10, 11]:
                cell.width = Inches(0.6)
            else:
                cell.width = Inches(0.475)

    i = 0
    get_merge_cells(table_residence_history, i, 0, 1)
    get_merge_cells(table_residence_history, i, 2, 5)
    get_merge_cells(table_residence_history, i, 6, 7)
    get_merge_cells(table_residence_history, i, 8, 9)

    current_row = table_residence_history.rows[i]

    current_row.cells[0].text = 'Previous Address'
    current_row.cells[6].text = 'Unit #'
    current_row.cells[10].text = 'Rent'
    current_row.cells[11].text = u"  □  "
    current_row.cells[12].text = 'Own'
    current_row.cells[13].text = u"  □  "

    set_cell_border(current_row.cells[2], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[8], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[11], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[13], bottom={"sz": 6, "color": "#000000", "val": "single"})

    i = 1
    get_merge_cells(table_residence_history, i, 6, 9)
    get_merge_cells(table_residence_history, i, 10, 13)

    current_row = table_residence_history.rows[i]

    current_row.cells[0].text = 'City'
    current_row.cells[2].text = 'State'
    current_row.cells[4].text = 'ZIP'
    current_row.cells[6].text = 'Monthly Payment or Rent $'

    set_cell_border(current_row.cells[1], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[3], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[5], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[10], bottom={"sz": 6, "color": "#000000", "val": "single"})

    i = 2
    get_merge_cells(table_residence_history, i, 0, 1)
    get_merge_cells(table_residence_history, i, 2, 5)
    get_merge_cells(table_residence_history, i, 6, 7)
    get_merge_cells(table_residence_history, i, 8, 9)
    get_merge_cells(table_residence_history, i, 10, 11)
    get_merge_cells(table_residence_history, i, 12, 13)

    current_row = table_residence_history.rows[i]

    current_row.cells[0].text = 'Dates of Residence'
    current_row.cells[2].text = f'__/__/____ to __/__/____'
    current_row.cells[6].text = 'Present Landlord'
    current_row.cells[10].text = 'Landlord Phone#'

    set_cell_border(current_row.cells[2], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[8], bottom={"sz": 6, "color": "#000000", "val": "single"})
    set_cell_border(current_row.cells[12], bottom={"sz": 6, "color": "#000000", "val": "single"})

    i = 3
    get_merge_cells(table_residence_history, i, 0, 1)
    get_merge_cells(table_residence_history, i, 2, len(table_residence_history.columns) - 1)

    current_row = table_residence_history.rows[i]

    current_row.cells[0].text = 'Reason for moving out'

    set_cell_border(current_row.cells[2], bottom={"sz": 6, "color": "#000000", "val": "single"})


def get_table_employment_information():
    table_employment_information = document.add_table(rows=5, cols=8)

    for row in table_employment_information.rows:
        for cell in row.cells:
            cell.width = Inches(0.7)

    for row in table_employment_information.rows:
        row.cells[0].width = Inches(1.4)
        row.cells[1].width = Inches(2.1)

    header_cell = get_merge_cells(table_employment_information, 0, 0, 7)
    header_cell.text = 'employment information'.upper()
    table_header_color = parse_xml(r'<w:shd {} w:fill="#1f3864"/>'.format(nsdecls('w')))
    header_cell._tc.get_or_add_tcPr().append(table_header_color)

    second_row = table_employment_information.rows[1]

    second_row.cells[0].text = 'Current Employer'

    set_cell_border(second_row.cells[1], bottom={"sz": 6, "color": "#000000", "val": "single"}, )

    get_merge_cells(table_employment_information, 1, 2, 3)
    second_row.cells[2].text = 'Position/Title'

    get_merge_cells(table_employment_information, 1, 4, 7)
    set_cell_border(second_row.cells[4], bottom={"sz": 6, "color": "#000000", "val": "single"}, )

    third_row = table_employment_information.rows[2]

    third_row.cells[0].text = 'Supervisor'

    set_cell_border(third_row.cells[1], bottom={"sz": 6, "color": "#000000", "val": "single"}, )

    get_merge_cells(table_employment_information, 2, 2, 3)
    third_row.cells[2].text = 'Phone #'

    get_merge_cells(table_employment_information, 2, 4, 7)
    set_cell_border(third_row.cells[4], bottom={"sz": 6, "color": "#000000", "val": "single"}, )

    fourth_row = table_employment_information.rows[3]

    fourth_row.cells[0].text = 'Address'
    set_cell_border(fourth_row.cells[1], bottom={"sz": 6, "color": "#000000", "val": "single"}, )
    fourth_row.cells[2].text = 'City'
    set_cell_border(fourth_row.cells[3], bottom={"sz": 6, "color": "#000000", "val": "single"}, )
    fourth_row.cells[4].text = 'State'
    set_cell_border(fourth_row.cells[5], bottom={"sz": 6, "color": "#000000", "val": "single"}, )
    fourth_row.cells[6].text = 'ZIP'
    set_cell_border(fourth_row.cells[7], bottom={"sz": 6, "color": "#000000", "val": "single"}, )

    fifth_row = table_employment_information.rows[4]

    fifth_row.cells[0].text = 'Dates of Employment'

    set_cell_border(fifth_row.cells[1], bottom={"sz": 6, "color": "#000000", "val": "single"}, )

    get_merge_cells(table_employment_information, 4, 2, 3)
    fifth_row.cells[2].text = 'Monthly Income $'

    get_merge_cells(table_employment_information, 4, 4, 7)
    set_cell_border(fifth_row.cells[4], bottom={"sz": 6, "color": "#000000", "val": "single"}, )


def get_table_employment_information_no_header():
    table_employment_information = document.add_table(rows=4, cols=8)

    for row in table_employment_information.rows:
        for cell in row.cells:
            cell.width = Inches(0.7)

    for row in table_employment_information.rows:
        row.cells[0].width = Inches(1.4)
        row.cells[1].width = Inches(2.1)

    second_row = table_employment_information.rows[0]

    second_row.cells[0].text = 'Previous Employer'

    set_cell_border(second_row.cells[1], bottom={"sz": 6, "color": "#000000", "val": "single"}, )

    get_merge_cells(table_employment_information, 0, 2, 3)
    second_row.cells[2].text = 'Position/Title'

    get_merge_cells(table_employment_information, 0, 4, 7)
    set_cell_border(second_row.cells[4], bottom={"sz": 6, "color": "#000000", "val": "single"}, )

    third_row = table_employment_information.rows[1]

    third_row.cells[0].text = 'Supervisor'

    set_cell_border(third_row.cells[1], bottom={"sz": 6, "color": "#000000", "val": "single"}, )

    get_merge_cells(table_employment_information, 1, 2, 3)
    third_row.cells[2].text = 'Phone #'

    get_merge_cells(table_employment_information, 1, 4, 7)
    set_cell_border(third_row.cells[4], bottom={"sz": 6, "color": "#000000", "val": "single"}, )

    fourth_row = table_employment_information.rows[2]

    fourth_row.cells[0].text = 'Address'
    set_cell_border(fourth_row.cells[1], bottom={"sz": 6, "color": "#000000", "val": "single"}, )
    fourth_row.cells[2].text = 'City'
    set_cell_border(fourth_row.cells[3], bottom={"sz": 6, "color": "#000000", "val": "single"}, )
    fourth_row.cells[4].text = 'State'
    set_cell_border(fourth_row.cells[5], bottom={"sz": 6, "color": "#000000", "val": "single"}, )
    fourth_row.cells[6].text = 'ZIP'
    set_cell_border(fourth_row.cells[7], bottom={"sz": 6, "color": "#000000", "val": "single"}, )

    fifth_row = table_employment_information.rows[3]

    fifth_row.cells[0].text = 'Dates of Employment'

    set_cell_border(fifth_row.cells[1], bottom={"sz": 6, "color": "#000000", "val": "single"}, )

    get_merge_cells(table_employment_information, 3, 2, 3)
    fifth_row.cells[2].text = 'Monthly Income $'

    get_merge_cells(table_employment_information, 3, 4, 7)
    set_cell_border(fifth_row.cells[4], bottom={"sz": 6, "color": "#000000", "val": "single"}, )


def get_table_additional_income():
    table_additional_income = document.add_table(rows=3, cols=7)
    header_cell = get_merge_cells(table_additional_income, 0, 0, 6)
    header_cell.text = 'ADDITIONAL_INCOME'
    table_header_color = parse_xml(r'<w:shd {} w:fill="#1f3864"/>'.format(nsdecls('w')))
    header_cell._tc.get_or_add_tcPr().append(table_header_color)

    table_additional_income.autofit = False

    for row_number in range(1,3):
        row = table_additional_income.rows[row_number].cells
        row[0].text = 'Source of Income'
        row_1_2 = get_merge_cells(table_additional_income, row_number, 1, 2)
        set_cell_border(row_1_2, bottom={"sz": 6, "color": "#000000", "val": "single"})
        row[3].text = 'Amount $'
        set_cell_border(row[4], bottom={"sz": 6, "color": "#000000", "val": "single"})
        row[5].text = 'Proof of Income'
        row[6].text = f'YES {CHECK_BOX} NO {CHECK_BOX}'

    paragraph = document.add_paragraph()
    paragraph.aligmnet = WD_TAB_ALIGNMENT
    paragraph_font = paragraph.style.font
    paragraph_font.name = 'Calibri'
    paragraph_font.size = Pt(10)


def get_table_vehicle_information():
    vehicle_information_text = ['Make','','Model','','Year','','State','','Plate #']
    table_vehicle_information = document.add_table(rows=3, cols=10)
    header_cell = get_merge_cells(table_vehicle_information, 0, 0, 9)
    header_cell.text = 'VEHICLE INFORMATION'
    table_header_color = parse_xml(r'<w:shd {} w:fill="#1f3864"/>'.format(nsdecls('w')))
    header_cell._tc.get_or_add_tcPr().append(table_header_color)

    table_vehicle_information.autofit = False

    for row_number in range(1,3):
        row = table_vehicle_information.rows[row_number].cells
        for cell in range(10):
            if cell%2 == 0:
                row[cell].text = vehicle_information_text[cell]
                row[cell].width = Inches(0.55)
            else:
                row[cell].width = Inches(1)
                set_cell_border(row[cell], bottom={"sz": 6, "color": "#000000", "val": "single"})

    paragraph = document.add_paragraph()
    paragraph.aligmnet = WD_TAB_ALIGNMENT
    paragraph_font = paragraph.style.font
    paragraph_font.name = 'Calibri'
    paragraph_font.size = Pt(10)


def get_table_pets_information():
    table_pets_information = document.add_table(rows=2, cols=13)
    header_cell = get_merge_cells(table_pets_information, 0, 0, 12)
    header_cell.text = 'VEHICLE INFORMATION'
    table_header_color = parse_xml(r'<w:shd {} w:fill="#1f3864"/>'.format(nsdecls('w')))
    header_cell._tc.get_or_add_tcPr().append(table_header_color)

    table_pets_information.autofit = False

    row = table_pets_information.rows[1].cells
    row[0].text = 'Pet(s)'
    row_1_2 = get_merge_cells(table_pets_information, 1, 1, 2)
    row_1_2.text = f'YES {CHECK_BOX} NO {CHECK_BOX}'
    row_3_4 = get_merge_cells(table_pets_information, 1, 3, 4)
    row_3_4.text = 'Numbers of Pets'
    set_cell_border(row[5], bottom={"sz": 6, "color": "#000000", "val": "single"})
    row[6].text = 'Type'
    row_7_8 = get_merge_cells(table_pets_information, 1, 7, 8)
    set_cell_border(row_7_8, bottom={"sz": 6, "color": "#000000", "val": "single"})
    row[9].text = 'Breed'
    row_10_12 = get_merge_cells(table_pets_information, 1, 10, 12)
    set_cell_border(row_10_12, bottom={"sz": 6, "color": "#000000", "val": "single"})

    paragraph = document.add_paragraph()
    paragraph.aligmnet = WD_TAB_ALIGNMENT
    paragraph_font = paragraph.style.font
    paragraph_font.name = 'Calibri'
    paragraph_font.size = Pt(10)


def get_table_personal_references():
    table_personal_references = document.add_table(rows=3, cols=8)
    header_cell = get_merge_cells(table_personal_references, 0, 0, 7)
    header_cell.text = 'PERSONAL REFERENCES'
    table_header_color = parse_xml(r'<w:shd {} w:fill="#1f3864"/>'.format(nsdecls('w')))
    header_cell._tc.get_or_add_tcPr().append(table_header_color)

    table_personal_references.autofit = False

    second_row = table_personal_references.rows[1].cells
    second_row[0].text = 'Full Name'
    second_1_3_row = get_merge_cells(table_personal_references, 1, 1, 3)
    set_cell_border(second_1_3_row, bottom={"sz": 6, "color": "#000000", "val": "single"})
    second_row[4].text = 'Address'
    second_5_7_row = get_merge_cells(table_personal_references, 1, 5, 7)
    set_cell_border(second_5_7_row, bottom={"sz": 6, "color": "#000000", "val": "single"})

    third_row = table_personal_references.rows[2].cells
    third_row[0].text = 'Relationship'
    third_1_3_row = get_merge_cells(table_personal_references, 2, 1, 3)
    set_cell_border(third_1_3_row, bottom={"sz": 6, "color": "#000000", "val": "single"})
    third_row[4].text = 'Phone #'
    third_5_7_row = get_merge_cells(table_personal_references, 2, 5, 7)
    set_cell_border(third_5_7_row, bottom={"sz": 6, "color": "#000000", "val": "single"})

    paragraph = document.add_paragraph()
    paragraph.aligmnet = WD_TAB_ALIGNMENT
    paragraph_font = paragraph.style.font
    paragraph_font.name = 'Calibri'
    paragraph_font.size = Pt(10)


def get_extra_table_personal_references():
    table_personal_references = document.add_table(rows=2, cols=8)
    table_personal_references.autofit = False

    second_row = table_personal_references.rows[0].cells
    second_row[0].text = 'Full Name'
    second_1_3_row = get_merge_cells(table_personal_references, 0, 1, 3)
    set_cell_border(second_1_3_row, bottom={"sz": 6, "color": "#000000", "val": "single"})
    second_row[4].text = 'Address'
    second_5_7_row = get_merge_cells(table_personal_references, 0, 5, 7)
    set_cell_border(second_5_7_row, bottom={"sz": 6, "color": "#000000", "val": "single"})

    third_row = table_personal_references.rows[1].cells
    third_row[0].text = 'Relationship'
    third_1_3_row = get_merge_cells(table_personal_references, 1, 1, 3)
    set_cell_border(third_1_3_row, bottom={"sz": 6, "color": "#000000", "val": "single"})
    third_row[4].text = 'Phone #'
    third_5_7_row = get_merge_cells(table_personal_references, 1, 5, 7)
    set_cell_border(third_5_7_row, bottom={"sz": 6, "color": "#000000", "val": "single"})

    paragraph = document.add_paragraph()
    paragraph.aligmnet = WD_TAB_ALIGNMENT
    paragraph_font = paragraph.style.font
    paragraph_font.name = 'Calibri'
    paragraph_font.size = Pt(10)


def get_table_additional_questions():
    questions = [
        'Have you ever broken a rental agreement?',
        'Have you ever been evicted or asked to move?',
        'Have you ever refused to pay the rent?',
        'Have you ever filed for bankruptcy?',
        'Have you ever been convicted of a crime?',
        'Are any of the occupants smokers?'
    ]
    table_additional_questions = document.add_table(rows=12, cols=12)
    header_cell = get_merge_cells(table_additional_questions, 0, 0, 11)
    header_cell.text = 'ADDITIONAL QUESTIONS'
    table_header_color = parse_xml(r'<w:shd {} w:fill="#1f3864"/>'.format(nsdecls('w')))
    header_cell._tc.get_or_add_tcPr().append(table_header_color)

    table_additional_questions.autofit = False

    for row_number in range(1,12):
        if row_number%2 != 0:
            row_0_4 = get_merge_cells(table_additional_questions, row_number, 0, 4)
            row_0_4.text = questions[row_number//2]
            row_5_6 = get_merge_cells(table_additional_questions, row_number, 5, 6)
            row_5_6.text = f'YES {CHECK_BOX} NO {CHECK_BOX}'
            row_7_11 = get_merge_cells(table_additional_questions, row_number, 7, 11)
            row_7_11.text = ''
        else:
            row_0_2 = get_merge_cells(table_additional_questions, row_number, 0, 2)
            row_0_2.text = 'If yes, provide explanation'
            row_3_11 = get_merge_cells(table_additional_questions, row_number, 3, 11)
            set_cell_border(row_3_11, bottom={"sz": 6, "color": "#000000", "val": "single"})


def get_additionals_questions_text():
    header_paragraph = document.add_paragraph()

    text = '''
    It is against a law to discriminate against any person in the terms, conditions, or privileges of sale or rental of 
    a dwelling, or in the provision of services or facilities in connection therewith, because of race, color, religion,
     sex, familial status, or national origin.

    By filling this form, applicant authorizes the verification of all information provided in this application,
     including Social Security Number, employment and income, credit history, previous and current rental history and 
     any other relevant information necessary for the Landlord to evaluate the application. 

    Applicant agrees that false or incomplete information filled in this application may result in a rejection 
    of this application and/or termination of a rental agreement.
    Non-refundable application fee: $20.00 
    '''
    header_paragraph.aligmnet = WD_TAB_ALIGNMENT
    header_paragraph_font = header_paragraph.style.font
    header_paragraph_font.name = 'Calibri'
    header_paragraph_font.size = Pt(12)
    header_paragraph.add_run(dedent(text)).bold = True


def get_table_sign():
    table_sign = document.add_table(4,3)

    first_row = table_sign.rows[0].cells
    first_row[0].text = ''
    set_cell_border(first_row[0], bottom={"sz": 6, "color": "#000000", "val": "single"})
    first_row[1].text = ''
    first_row[2].text = ''

    set_cell_border(first_row[2], bottom={"sz": 6, "color": "#000000", "val": "single"})

    second_row = table_sign.rows[1].cells
    aplicant_parg = second_row[0].add_paragraph('Applicant Signature')
    aplicant_parg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    second_row[1].text = ''
    landlord_parg = second_row[2].add_paragraph('Landlord Signature')
    landlord_parg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    third_row = table_sign.rows[2].cells
    third_row[0].text = ''
    set_cell_border(third_row [0], bottom={"sz": 6, "color": "#000000", "val": "single"})
    third_row[1].text = ''
    set_cell_border(third_row [2], bottom={"sz": 6, "color": "#000000", "val": "single"})

    fourth_row = table_sign.rows[3].cells
    first_date_parg = fourth_row[0].add_paragraph('Date')
    first_date_parg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fourth_row[1].text = ''
    first_second_date = fourth_row[2].add_paragraph('Date')
    first_second_date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph()
    paragraph.aligmnet = WD_TAB_ALIGNMENT
    paragraph_font = paragraph.style.font
    paragraph_font.name = 'Calibri'
    paragraph_font.size = Pt(10)

    # for row in range(4):
    #     table_sign.rows[row].cells[0].width = Inches(2.5)
    # table_sign.rows[0].cells[0].text = ''
    # signature = table_sign.rows[1].cells[0].add_paragraph('Landlord Signature')
    # signature.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # # table_sign.rows[1].cells[0].text = 'Landlord Signature'
    # table_sign.rows[2].cells[0].text = ''
    # # set_cell_border(table_sign.rows[2].cells[0], bottom={"sz": 6, "color": "#000000", "val": "single"}, )
    # table_sign.rows[3].cells[0].text = 'Date'


if __name__=='__main__':

    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    get_header()
    get_header_paragraph('test')
    get_table_applicant_information()
    get_table_additional_occupant()
    get_table_residence_history()
    document.add_paragraph()
    get_table_residence_history_no_header()
    document.add_paragraph()
    get_table_residence_history_no_header()
    document.add_page_break()
    get_table_employment_information()
    document.add_paragraph()
    get_table_employment_information_no_header()
    document.add_paragraph()
    get_table_employment_information_no_header()
    document.add_paragraph()
    get_table_additional_income()
    document.add_paragraph()
    get_table_vehicle_information()
    document.add_paragraph()
    get_table_pets_information()
    document.add_paragraph()
    get_table_personal_references()
    get_extra_table_personal_references()
    get_extra_table_personal_references()
    document.add_paragraph()
    get_table_additional_questions()
    get_additionals_questions_text()
    get_table_sign()

    for table in document.tables:
        set_table_alignment(table)

    document.save('test_doc.docx')

